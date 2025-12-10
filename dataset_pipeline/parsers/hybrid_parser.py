from __future__ import annotations

import logging
from pathlib import Path
from typing import Optional

from dataset_pipeline.core.schemas import ParsedDocument
from dataset_pipeline.core.utils import merge_metadata, normalise_whitespace, safe_read_text

from .rule_based import RuleBasedParser
from .vlm_parser import VLMParser

logger = logging.getLogger(__name__)


class HybridParser:
    """Combines text extraction with OCR fallback for DOCX/PDF."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown"}

    def __init__(self, rule_parser: RuleBasedParser, vlm_parser: Optional[VLMParser] = None) -> None:
        self.rule_parser = rule_parser
        self.vlm_parser = vlm_parser

    def parse(self, path: Path, doc_type: str) -> Optional[ParsedDocument]:
        text = self._extract_text(path)
        if text:
            cleaned = normalise_whitespace(text)
            if not cleaned:
                text = ""
            else:
                metadata = {"parser_type": "hybrid", "format": path.suffix.lstrip(".") or path.suffix}
                if doc_type == "knowledge":
                    return ParsedDocument(
                        source=str(path),
                        doc_type="knowledge",
                        parser_type="hybrid",
                        format=path.suffix.lstrip(".") or "doc",
                        knowledge=cleaned,
                        metadata=metadata,
                    )
                doc = self.rule_parser.parse_text(cleaned, path, parser_type="hybrid", format_hint="txt")
                if doc:
                    merged_meta = merge_metadata(doc.metadata, metadata)
                    return doc.copy(update={"metadata": merged_meta, "parser_type": "hybrid"})

        if self.vlm_parser:
            logger.info("Hybrid parser falling back to VLM for %s", path)
            return self.vlm_parser.parse(path, doc_type=doc_type)
        return None

    def _extract_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".md", ".markdown"}:
            return safe_read_text(path)
        if suffix == ".docx":
            return self._extract_docx(path)
        if suffix == ".pdf":
            return self._extract_pdf(path)
        return ""

    def _extract_docx(self, path: Path) -> str:
        try:
            import docx2txt  # type: ignore
        except ImportError:
            try:
                from docx import Document  # type: ignore
            except ImportError as exc:  # pragma: no cover - optional dependency
                logger.warning("Neither docx2txt nor python-docx available to parse %s: %s", path, exc)
                return ""
            try:
                document = Document(str(path))
            except Exception as exc:  # pragma: no cover
                logger.warning("Failed to open DOCX %s: %s", path, exc)
                return ""
            return "\n".join(para.text for para in document.paragraphs if para.text)
        try:
            return docx2txt.process(path) or ""
        except Exception as exc:  # pragma: no cover
            logger.warning("docx2txt failed for %s: %s", path, exc)
            return ""

    def _extract_pdf(self, path: Path) -> str:
        try:
            from pdfminer.high_level import extract_text  # type: ignore
        except ImportError as exc:  # pragma: no cover - optional dependency
            logger.warning("pdfminer.six not installed; cannot parse %s: %s", path, exc)
            return ""
        try:
            return extract_text(str(path))
        except Exception as exc:  # pragma: no cover
            logger.warning("Failed to parse PDF %s: %s", path, exc)
            return ""
